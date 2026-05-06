"""
文档扫描模块

支持 PDF/DOCX/TXT/MD 多种文档格式的安全扫描，检测敏感词和隐私数据。
性能要求：<3 秒/百页

增强功能（2026-04-14）：
- 集成 ragshield-rules 规则库（518 敏感词 + 14 隐私正则）
- 使用统一 RulesLoader 加载规则
"""

import json
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple, Union

from ..utils.logger import get_logger
from ..utils.metrics import global_collector
from .rules_loader import RulesLoader

logger = get_logger(__name__)


@dataclass
class Issue:
    """扫描发现的问题"""
    rule_id: str
    rule_name: str
    level: str  # high|medium|low
    position: Tuple[int, int]  # (line, column)
    suggestion: str
    matched_text: str = ""
    rule_type: str = ""


@dataclass
class ScanResult:
    """扫描结果"""
    passed: bool
    score: int  # 0-100
    issues: List[Issue] = field(default_factory=list)
    duration_ms: int = 0
    doc_type: str = ""
    page_count: int = 0
    word_count: int = 0

    @property
    def is_valid(self) -> bool:
        """别名，兼容测试用例"""
        return self.passed


class DocScanner:
    """文档扫描器（增强版 - 使用 ragshield-rules）"""

    def __init__(self, rules_dir: Optional[str] = None):
        """
        初始化文档扫描器

        Args:
            rules_dir: 规则文件目录（可选，默认使用 ragshield-rules）
        """
        # 使用统一规则加载器
        self.loader = RulesLoader(rules_dir)

        # 加载敏感词和隐私规则
        self.sensitive_words = self.loader.load_sensitive_words()
        self.privacy_rules = self.loader.load_privacy_rules()

        # 预编译隐私正则
        self._privacy_patterns = []
        for rule in self.privacy_rules:
            patterns = rule.get("patterns", [])
            for pattern in patterns:
                try:
                    compiled = re.compile(pattern)
                    self._privacy_patterns.append((rule, compiled))
                except re.error as e:
                    logger.warning(f"隐私正则编译失败 {rule['rule_id']}: {e}")

        logger.info(f"DocScanner 初始化完成（ragshield-rules），加载 {len(self.sensitive_words)} 个敏感词，"
                   f"{len(self.privacy_rules)} 条隐私规则")

    def scan(
        self,
        content: Union[str, bytes],
        doc_type: str = "text",
        rules: Optional[List[str]] = None
    ) -> ScanResult:
        """
        扫描文档内容
        
        Args:
            content: 文档内容（字符串或字节）
            doc_type: 文档类型 (text|pdf|docx|md)
            rules: 启用的规则列表 (None 表示全部启用)
        
        Returns:
            ScanResult 扫描结果
        """
        start_time = time.perf_counter()
        
        # 解析文档
        text, page_count = self._parse_document(content, doc_type)
        
        # 执行扫描
        issues = []
        
        # 敏感词检测
        if not rules or "sensitive_word" in rules:
            issues.extend(self._check_sensitive_words(text))
        
        # 隐私数据检测
        if not rules or "privacy_data" in rules:
            issues.extend(self._check_privacy_data(text))
        
        # 计算评分
        score = self._calculate_score(issues)
        
        # 修复：有 issues 就返回 False
        passed = len(issues) == 0
        
        duration_ms = int((time.perf_counter() - start_time) * 1000)
        
        result = ScanResult(
            passed=passed,
            score=score,
            issues=issues,
            duration_ms=duration_ms,
            doc_type=doc_type,
            page_count=page_count,
            word_count=len(text.split())
        )
        
        logger.info(f"文档扫描完成：{doc_type}, {page_count}页，{len(issues)}个问题，{duration_ms}ms, 得分{score}")
        
        return result
    
    def _parse_document(self, content: Union[str, bytes], doc_type: str) -> Tuple[str, int]:
        """
        解析文档，提取纯文本
        
        Returns:
            (文本内容，页数)
        """
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        if doc_type == "text" or doc_type == "txt":
            return content, max(1, content.count('\n') // 30 + 1)
        
        elif doc_type == "md" or doc_type == "markdown":
            return content, max(1, content.count('\n') // 30 + 1)
        
        elif doc_type == "pdf":
            return self._parse_pdf(content)
        
        elif doc_type == "docx":
            return self._parse_docx(content)
        
        else:
            logger.warning(f"未知的文档类型：{doc_type}，按文本处理")
            return content, max(1, content.count('\n') // 30 + 1)
    
    def _parse_pdf(self, content: Union[str, bytes]) -> Tuple[str, int]:
        """解析 PDF 文件"""
        try:
            import pdfplumber
            
            if isinstance(content, str):
                # 如果是字符串，假设是文件路径
                with pdfplumber.open(content) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                    return text, len(pdf.pages)
            else:
                # 字节内容需要临时文件
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
                    f.write(content)
                    temp_path = f.name
                
                try:
                    with pdfplumber.open(temp_path) as pdf:
                        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                        return text, len(pdf.pages)
                finally:
                    Path(temp_path).unlink()
        
        except ImportError:
            logger.error("pdfplumber 未安装，无法解析 PDF")
            return str(content), 1
        except Exception as e:
            logger.error(f"PDF 解析失败：{e}")
            return str(content), 1
    
    def _parse_docx(self, content: Union[str, bytes]) -> Tuple[str, int]:
        """解析 DOCX 文件"""
        try:
            import docx
            
            if isinstance(content, str):
                doc = docx.Document(content)
            else:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                    f.write(content)
                    temp_path = f.name
                
                try:
                    doc = docx.Document(temp_path)
                finally:
                    Path(temp_path).unlink()
            
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            # DOCX 没有明确的页数概念，估算
            page_count = max(1, len(text) // 1500 + 1)
            return text, page_count
        
        except ImportError:
            logger.error("python-docx 未安装，无法解析 DOCX")
            return str(content), 1
        except Exception as e:
            logger.error(f"DOCX 解析失败：{e}")
            return str(content), 1
    
    def _check_sensitive_words(self, text: str) -> List[Issue]:
        """检测敏感词（使用 ragshield-rules 518 敏感词）"""
        issues = []
        lines = text.split('\n')

        for word in self.sensitive_words:
            for line_num, line in enumerate(lines, 1):
                # 简单字符串匹配（后续可优化为 Aho-Corasick）
                if word.lower() in line.lower():
                    col = line.lower().index(word.lower()) + 1
                    issues.append(Issue(
                        rule_id="SENS001",
                        rule_name="敏感词",
                        level="medium",
                        position=(line_num, col),
                        suggestion="文档包含敏感内容，建议脱敏处理",
                        matched_text=word,
                        rule_type="keyword"
                    ))

        return issues

    def _check_privacy_data(self, text: str) -> List[Issue]:
        """检测隐私数据（使用预编译的隐私正则）"""
        issues = []
        lines = text.split('\n')

        for rule, compiled in self._privacy_patterns:
            for line_num, line in enumerate(lines, 1):
                for match in compiled.finditer(line):
                    issues.append(Issue(
                        rule_id=rule["rule_id"],
                        rule_name=rule["rule_name"],
                        level=rule.get("level", "high"),
                        position=(line_num, match.start() + 1),
                        suggestion=rule.get("suggestion", "检测到隐私数据，建议脱敏处理"),
                        matched_text=match.group(),
                        rule_type="regex"
                    ))

        return issues
    
    def _calculate_score(self, issues: List[Issue]) -> int:
        """
        计算安全评分
        
        算法：
        - 基础分 100 分
        - high 级别问题：-15 分/个
        - medium 级别问题：-8 分/个
        - low 级别问题：-3 分/个
        - 最低 0 分
        """
        score = 100
        
        for issue in issues:
            if issue.level == "high":
                score -= 15
            elif issue.level == "medium":
                score -= 8
            elif issue.level == "low":
                score -= 3
        
        return max(0, score)


# 便捷函数
_scanner_instance: Optional[DocScanner] = None


def doc_scanner(
    content: Union[str, bytes],
    doc_type: str = "text",
    rules: Optional[List[str]] = None
) -> ScanResult:
    """
    文档扫描便捷函数
    
    Args:
        content: 文档内容
        doc_type: 文档类型 (text|pdf|docx|md)
        rules: 启用的规则列表
    
    Returns:
        ScanResult 扫描结果
    """
    global _scanner_instance
    if _scanner_instance is None:
        _scanner_instance = DocScanner()
    
    with global_collector.measure("doc_scanner"):
        return _scanner_instance.scan(content, doc_type, rules)
